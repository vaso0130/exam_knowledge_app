import os
import requests
from typing import Tuple
from charset_normalizer import from_path
import tempfile

# Google Cloud Vision OCR
try:
    from google.cloud import vision
    from google.oauth2 import service_account
except ImportError:
    vision = None
    service_account = None

# PDF 轉圖片
try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

# 文件處理相關
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from bs4 import BeautifulSoup  
except ImportError:
    BeautifulSoup = None


class GoogleVisionOCR:
    """Google Cloud Vision OCR 處理器"""
    
    def __init__(self, credentials_path: str = "google_credentials.json"):
        """初始化 Google Vision OCR 客戶端"""
        self.client = None
        self.credentials_path = credentials_path
        
        if vision is None:
            print("警告：google-cloud-vision 未安裝，OCR 功能將不可用")
            return
        
        try:
            if os.path.exists(credentials_path):
                os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = credentials_path
                credentials = service_account.Credentials.from_service_account_file(credentials_path)
                self.client = vision.ImageAnnotatorClient(credentials=credentials)
            else:
                print(f"警告：找不到憑證檔案 {credentials_path}，OCR 功能將不可用")
        except Exception as e:
            print(f"警告：初始化 Google Vision OCR 失敗: {e}")
    
    def extract_text_from_image(self, image_path: str) -> str:
        """從圖片中提取文字，並返回帶有幾何座標的詞彙列表"""
        if self.client is None:
            raise ValueError("Google Vision OCR 客戶端未初始化")

        try:
            with open(image_path, 'rb') as image_file:
                content = image_file.read()

            image = vision.Image(content=content)
            response = self.client.document_text_detection(image=image)

            if response.error.message:
                raise Exception(f'Google Vision API 錯誤: {response.error.message}')

            words_with_coords = []
            if response.full_text_annotation:
                for page in response.full_text_annotation.pages:
                    for block in page.blocks:
                        for paragraph in block.paragraphs:
                            for word in paragraph.words:
                                word_text = ''.join([symbol.text for symbol in word.symbols])
                                vertices = word.bounding_box.vertices
                                words_with_coords.append({
                                    'text': word_text,
                                    'x0': vertices[0].x,
                                    'top': vertices[0].y,
                                })
            return FileProcessor._reconstruct_text_from_words(words_with_coords)

        except Exception as e:
            raise ValueError(f"OCR 處理失敗: {e}")
    
    def extract_text_from_pdf_pages(self, pdf_path: str) -> str:
        """從 PDF 頁面中提取文字（通過轉換為圖片）"""
        if convert_from_path is None:
            raise ImportError("請安裝 pdf2image 套件：pip install pdf2image")
        
        if self.client is None:
            raise ValueError("Google Vision OCR 客戶端未初始化")
        
        try:
            pages = convert_from_path(pdf_path, dpi=300)
            extracted_text = []
            
            for page in pages:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                    page.save(temp_file.name, 'PNG')
                    temp_image_path = temp_file.name
                
                try:
                    page_text = self.extract_text_from_image(temp_image_path)
                    if page_text.strip():
                        extracted_text.append(page_text)
                finally:
                    if os.path.exists(temp_image_path):
                        os.unlink(temp_image_path)
            
            return "\n\n".join(extracted_text)
            
        except Exception as e:
            raise ValueError(f"PDF OCR 處理失敗: {e}")


class FileProcessor:
    """檔案處理器，支援多種格式"""
    
    def __init__(self):
        """初始化檔案處理器"""
        self.ocr = GoogleVisionOCR()

    @staticmethod
    def _reconstruct_text_from_words(words: list) -> str:
        """
        [共用函式] 根據帶座標的詞彙列表，重建包含縮排的完整文字。
        此函式可同時被 pdfplumber 和 Google Vision OCR 的結果使用。
        """
        if not words:
            return ""

        words.sort(key=lambda w: (w['top'], w['x0']))
        
        lines = []
        if not words:
            return ""
        current_line = [words[0]]
        line_tolerance = 2
        for i in range(1, len(words)):
            last_word_in_line = sorted(current_line, key=lambda w: w['x0'])[-1]
            current_word = words[i]
            
            if abs(current_word['top'] - last_word_in_line['top']) <= line_tolerance:
                current_line.append(current_word)
            else:
                lines.append(sorted(current_line, key=lambda w: w['x0']))
                current_line = [current_word]
        lines.append(sorted(current_line, key=lambda w: w['x0']))

        if not lines:
            return ""

        base_indent = min(line[0]['x0'] for line in lines if line)
        full_text_content = []
        
        for line_words in lines:
            if not line_words: continue
            
            first_word_indent = line_words[0]['x0']
            indent_pixels = first_word_indent - base_indent
            char_width = 5 
            num_spaces = int(round(indent_pixels / char_width))
            
            line_text = ' '.join(w['text'] for w in line_words)
            reconstructed_line = ' ' * num_spaces + line_text
            full_text_content.append(reconstructed_line)
            
        return "\n".join(full_text_content)

    @staticmethod
    def preprocess_pseudocode(text: str) -> str:
        if not text:
            return ""
        cleaned_lines = []
        for line in text.splitlines():
            line = line.replace('\u3000', '  ').replace('\xa0', ' ')
            line = line.replace('←', '->').replace('→', '->')
            line = line.replace('\x00', '')
            cleaned_lines.append(line)
        return "\n".join(cleaned_lines)
    
    @staticmethod
    def read_text_file(file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return FileProcessor.preprocess_pseudocode(f.read())
        except UnicodeDecodeError:
            pass
        try:
            detected = from_path(file_path).best()
            if detected: return FileProcessor.preprocess_pseudocode(str(detected))
        except Exception:
            pass
        encodings = ['big5', 'gbk', 'cp1252', 'latin1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    return FileProcessor.preprocess_pseudocode(f.read())
            except Exception: continue
        raise ValueError(f"無法讀取檔案 {file_path}，編碼格式不支援")

    def read_pdf_file(self, file_path: str) -> str:
        """
        讀取PDF檔案，優先使用幾何分析，若失敗則回退到 OCR。
        """
        if pdfplumber is None:
            raise ImportError("請安裝 pdfplumber 套件：pip install pdfplumber")

        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                all_words = []
                for page in pdf.pages:
                    all_words.extend(page.extract_words(use_text_flow=True, x_tolerance=2))
                text = self._reconstruct_text_from_words(all_words)
        except Exception as e:
            print(f"使用 pdfplumber 進行幾何分析失敗: {e}。嘗試後備方案。")
            text = ""

        if len(text.strip()) < 50:
            print("幾何分析未提取到足夠文字，啟動 OCR 掃描備用方案...")
            try:
                text = self.ocr.extract_text_from_pdf_pages(file_path)
            except Exception as ocr_error:
                raise ValueError(f"PDF 檔案讀取失敗：主要方法和 OCR 備用方案均失敗。({ocr_error})")

        return self.preprocess_pseudocode(text)
    
    def read_image_file(self, file_path: str) -> str:
        """讀取圖片檔案並進行 OCR，使用幾何分析重建排版"""
        try:
            text = self.ocr.extract_text_from_image(file_path)
            return self.preprocess_pseudocode(text)
        except Exception as e:
            raise ValueError(f"無法讀取圖片檔案 {file_path}: {e}")
    
    @staticmethod 
    def read_docx_file(file_path: str) -> str:
        """讀取Word檔案"""
        if Document is None:
            raise ImportError("請安裝 python-docx 套件：pip install python-docx")
        
        try:
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return FileProcessor.preprocess_pseudocode('\n'.join(text))
        except Exception as e:
            raise ValueError(f"無法讀取Word檔案 {file_path}: {e}")
    
    @staticmethod
    def read_html_file(file_path: str) -> str:
        """讀取HTML檔案"""
        if BeautifulSoup is None:
            raise ImportError("請安裝 beautifulsoup4 套件：pip install beautifulsoup4")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            # 移除 script 和 style 標籤
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator='\n').strip()
            return FileProcessor.preprocess_pseudocode(text)
        except Exception as e:
            raise ValueError(f"無法讀取HTML檔案 {file_path}: {e}")
    
    @staticmethod
    def fetch_url_content_sync(url: str) -> str:
        """同步版本的 URL 內容獲取 - 用於 Flask 路由"""
        try:
            # 在新的線程中創建事件循環
            import asyncio
            import threading
            
            result_container = {'result': None, 'error': None}
            
            def run_async():
                try:
                    # 創建新的事件循環
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    try:
                        result = loop.run_until_complete(FileProcessor._fetch_url_async(url))
                        result_container['result'] = result
                    finally:
                        loop.close()
                except Exception as e:
                    result_container['error'] = e
            
            # 在獨立線程中運行異步代碼
            thread = threading.Thread(target=run_async)
            thread.start()
            thread.join(timeout=60)  # 60秒超時
            
            if result_container['error']:
                raise result_container['error']
            
            if result_container['result']:
                return result_container['result']
            else:
                # 如果沒有結果，回退到傳統方法
                return FileProcessor._fetch_url_fallback(url)
                
        except Exception as e:
            print(f"Playwright 抓取失敗，使用傳統方法: {e}")
            return FileProcessor._fetch_url_fallback(url)
    
    @staticmethod
    def _fetch_url_fallback(url: str) -> str:
        """傳統方法的 URL 抓取回退方案"""
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # 移除不需要的元素
            for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
                script.decompose()
            
            # 嘗試找到主要內容
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            if main_content:
                text = main_content.get_text(separator='\n').strip()
            else:
                text = soup.get_text(separator='\n').strip()
            
            # 清理文字
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            cleaned_text = '\n'.join(lines)
            
            # 限制內容長度，避免巨量文字消耗配額
            max_length = 100000  # 100KB 文字限制
            if len(cleaned_text) > max_length:
                cleaned_text = cleaned_text[:max_length] + "\n\n... (內容過長，已截斷)"
            
            return FileProcessor.preprocess_pseudocode(cleaned_text)
            
        except ImportError:
            response_text = response.text
            # 同樣限制長度
            max_length = 100000
            if len(response_text) > max_length:
                response_text = response_text[:max_length] + "\n\n... (內容過長，已截斷)"
            return FileProcessor.preprocess_pseudocode(response_text)
    
    @staticmethod
    async def _fetch_url_async(url: str) -> str:
        """異步獲取 URL 內容的內部方法"""
        try:
            from .playwright_scraper import scrape_single_page
            
            # 執行異步爬取
            result = await scrape_single_page(url, headless=True)
            
            if result['status'] == 'success':
                # 建構完整內容，包含表格
                content_parts = []
                
                # 添加標題
                if result['title']:
                    content_parts.append(f"# {result['title']}\n")
                
                # 添加表格（優先處理）
                if result['tables']:
                    content_parts.append("## 表格資訊\n")
                    for i, table in enumerate(result['tables']):
                        if table['markdown']:
                            content_parts.append(f"### 表格 {i+1}\n{table['markdown']}\n")
                
                # 添加主要文字內容
                if result['text_content']:
                    content_parts.append("## 主要內容\n")
                    content_parts.append(result['text_content'])
                
                # 添加圖片資訊 - 只用原始連結，不做 base64 處理
                if result['images']:
                    content_parts.append("\n## 圖片資訊\n")
                    image_count = 0
                    for i, img in enumerate(result['images']):
                        if image_count >= 5:
                            break
                        img_src = img['src']
                        img_alt = img['alt'] or img['title'] or f"圖片 {i+1}"
                        content_parts.append(f"![{img_alt}]({img_src})")
                        image_count += 1
                
                full_content = '\n\n'.join(content_parts)
                
                # 限制最終內容大小，避免過度消耗 API 配額
                max_length = 100000  # 100KB 文字限制
                if len(full_content) > max_length:
                    full_content = full_content[:max_length] + "\n\n... (內容過長，已截斷)"
                
                if len(full_content.strip()) > 50:  # 確保有足夠內容
                    return FileProcessor.preprocess_pseudocode(full_content)
            
        except ImportError:
            pass  # Playwright 未安裝，回退到傳統方法
        except Exception as e:
            print(f"Playwright 爬取失敗，回退到傳統方法: {e}")
        
        # 回退到傳統的 requests + BeautifulSoup 方法
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(response.content, 'html.parser')
            # 移除 script 和 style 標籤
            for script in soup(["script", "style"]):
                script.decompose()
            
            text_content = soup.get_text(separator='\n').strip()
            
            # 限制內容長度，避免巨量文字消耗配額
            max_length = 100000  # 100KB 文字限制
            if len(text_content) > max_length:
                text_content = text_content[:max_length] + "\n\n... (內容過長，已截斷)"
            
            return FileProcessor.preprocess_pseudocode(text_content)
        except ImportError:
            response_text = response.text
            # 同樣限制長度
            max_length = 100000
            if len(response_text) > max_length:
                response_text = response_text[:max_length] + "\n\n... (內容過長，已截斷)"
            return FileProcessor.preprocess_pseudocode(response_text)

    # 已移除 base64 相關圖片處理，所有圖片只用原始連結

    @staticmethod
    async def fetch_url_content(url: str) -> str:
        """從URL獲取內容 - 使用 Playwright 進階爬取"""
        return await FileProcessor._fetch_url_async(url)
    
    @classmethod
    def process_input(cls, input_data: str) -> Tuple[str, str]:
        """
        處理輸入資料（檔案路徑、URL或純文字）
        返回：(處理後的文字內容, 輸入類型)
        """
        input_data = input_data.strip()
        
        # 建立處理器實例
        processor = cls()
        
        # 檢查是否為URL
        if input_data.startswith(('http://', 'https://')):
            try:
                content = cls.fetch_url_content_sync(input_data)
                return content, 'url'
            except Exception as e:
                raise ValueError(f"URL處理失敗: {e}")
        
        # 檢查是否為檔案路徑
        if os.path.exists(input_data):
            file_ext = os.path.splitext(input_data)[1].lower()
            
            if file_ext == '.txt':
                content = cls.read_text_file(input_data)
                return content, 'txt'
            elif file_ext == '.pdf':
                content = processor.read_pdf_file(input_data)
                return content, 'pdf'
            elif file_ext in ['.docx', '.doc']:
                content = cls.read_docx_file(input_data)
                return content, 'docx'
            elif file_ext in ['.html', '.htm']:
                content = cls.read_html_file(input_data)
                return content, 'html'
            elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp']:
                # 圖片檔案，使用 OCR 提取文字
                content = processor.read_image_file(input_data)
                return content, 'image'
            else:
                # 嘗試當作純文字檔案讀取
                try:
                    content = cls.read_text_file(input_data)
                    return content, 'txt'
                except Exception:
                    raise ValueError(f"不支援的檔案格式: {file_ext}")
        
        # 當作純文字處理
        return input_data, 'text'
    
    @staticmethod
    def save_markdown(content: str, file_path: str) -> None:
        """儲存Markdown檔案"""
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    @staticmethod
    def generate_markdown_filename(subject: str, hash_str: str) -> str:
        """生成Markdown檔案名稱"""
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{subject}/{timestamp}_{hash_str}.md"

class ContentValidator:
    """內容驗證器"""
    
    @staticmethod
    def is_valid_text(text: str, min_length: int = 10) -> bool:
        """檢查文字是否有效"""
        return text and len(text.strip()) >= min_length
    
    @staticmethod
    def clean_text(text: str) -> str:
        """清理文字內容"""
        if not text:
            return ""
        
        # 移除多餘的空白行
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            if line.strip():  # 只保留非空行
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    @staticmethod
    def truncate_text(text: str, max_length: int = 10000) -> str:
        """截斷過長的文字"""
        if len(text) <= max_length:
            return text
        return text[:max_length] + "..."